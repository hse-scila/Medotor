"""
Module for processing files of various formats
PDF, TXT, DOC, DOCX support for RAG system
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import aiofiles
import tempfile

# Logging configuration (before importing libraries)
logger = logging.getLogger(__name__)

# Imports for file processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
    logger.info(f"PyPDF2 available, version: {PyPDF2.__version__}")
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed: pip install PyPDF2")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed: pip install python-docx")

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("python-magic not installed: pip install python-magic")

class FileProcessor:
    """Class for processing files of various formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.doc': self._process_docx,  # Use docx for .doc
            '.docx': self._process_docx
        }
        
        # Create temporary folder for files
        self.temp_dir = Path("temp/uploads")
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported formats"""
        return list(self.supported_formats.keys())
    
    def is_supported_format(self, filename: str) -> bool:
        """Check file format support"""
        ext = Path(filename).suffix.lower()
        return ext in self.supported_formats
    
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file to temporary folder"""
        file_path = self.temp_dir / filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        logger.info(f"File saved: {file_path}")
        return str(file_path)
    
    async def process_file(self, file_path: str, filename: str, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None) -> Dict[str, Any]:
        """Process file and extract text"""
        try:
            file_path = Path(file_path)
            ext = file_path.suffix.lower()
            
            if ext not in self.supported_formats:
                raise ValueError(f"Неподдерживаемый формат файла: {ext}")
            
            # Determine MIME type if available
            mime_type = "unknown"
            if MAGIC_AVAILABLE:
                try:
                    mime_type = magic.from_file(str(file_path), mime=True)
                except:
                    pass
            
            # Extract text
            processor = self.supported_formats[ext]
            text_content = await processor(str(file_path))
            
            # Split into chunks for better RAG processing
            chunks = self._split_text_into_chunks(text_content, chunk_size=chunk_size, overlap=chunk_overlap)
            
            result = {
                "filename": filename,
                "file_path": str(file_path),
                "format": ext,
                "mime_type": mime_type,
                "text_content": text_content,
                "chunks": chunks,
                "chunks_count": len(chunks),
                "total_length": len(text_content),
                "status": "success"
            }
            
            logger.info(f"File processed: {filename} ({len(chunks)} chunks)")
            return result
            
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            logger.error(f"Error processing file {filename}: {error_type}: {error_message}")
            import traceback
            logger.error(f"Traceback:\n{traceback.format_exc()}")
            return {
                "filename": filename,
                "status": "error",
                "error": f"{error_type}: {error_message}"
            }
    
    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 не установлен. Установите: pip install PyPDF2")
        
        text_content = ""
        file_path_obj = Path(file_path)
        
        # Check file existence
        if not file_path_obj.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        # Check file size
        file_size = file_path_obj.stat().st_size
        logger.info(f"Processing PDF: {file_path}, size: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("PDF файл пуст (размер 0 байт)")
        
        try:
            with open(file_path, 'rb') as file:
                # Check first bytes of file (PDF should start with %PDF)
                first_bytes = file.read(4)
                file.seek(0)
                if not first_bytes.startswith(b'%PDF'):
                    raise ValueError(f"Файл не является корректным PDF (не найден заголовок %PDF). Первые байты: {first_bytes}")
                
                logger.info(f"PDF header is correct, starting reading...")
                pdf_reader = PyPDF2.PdfReader(file)
                
                logger.info(f"PDF reader created, pages: {len(pdf_reader.pages)}")
                
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF файл не содержит страниц или поврежден")
                
                pages_with_text = 0
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content += page_text + "\n"
                            pages_with_text += 1
                            logger.debug(f"Page {page_num + 1}: extracted {len(page_text)} characters")
                        else:
                            logger.warning(f"Page {page_num + 1}: text is empty or missing")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {type(e).__name__}: {e}")
                        continue
                
                logger.info(f"Extracted text: {len(text_content)} characters from {pages_with_text} pages out of {len(pdf_reader.pages)}")
                
                if not text_content.strip():
                    raise ValueError("Не удалось извлечь текст из PDF. Возможно, это сканированный документ (изображения). Используйте OCR или конвертируйте PDF в изображения.")
                
                return text_content.strip()
        except PyPDF2.errors.PdfReadError as e:
            error_msg = f"Ошибка чтения PDF файла: {str(e)}. Файл может быть поврежден, зашифрован или иметь несовместимый формат."
            logger.error(error_msg)
            raise ValueError(error_msg)
        except ValueError as e:
            # Re-raise ValueError as is
            logger.error(f"ValueError when processing PDF: {e}")
            raise
        except Exception as e:
            error_msg = f"Неожиданная ошибка обработки PDF: {type(e).__name__}: {str(e)}"
            logger.error(f"Error processing PDF {file_path}: {error_msg}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise ValueError(error_msg)
    
    async def _process_txt(self, file_path: str) -> str:
        """Process TXT file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        return content.strip()
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOC/DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен")
        
        doc = Document(file_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"
        
        return text_content.strip()
    
    def _split_text_into_chunks(self, text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None) -> List[str]:
        """Разбиение текста на чанки для лучшей обработки в RAG"""
        if not text:
            return []
        
        # Use settings from configuration if parameters not specified
        if chunk_size is None or overlap is None:
            try:
                from config import get_config
                config = get_config()
                if chunk_size is None:
                    chunk_size = config.rag.chunk_size
                if overlap is None:
                    overlap = config.rag.chunk_overlap
            except Exception:
                chunk_size = chunk_size or 100
                overlap = overlap or 40
        
        # Maximum number of chunks to protect against very large documents
        max_chunks = None
        try:
            from config import get_config
            config = get_config()
            max_chunks = config.rag.max_chunks_per_document
        except Exception:
            max_chunks = None

        # For very large texts use fast character mode
        # to avoid huge words list and hangs
        if len(text) > 2_000_000:
            chunks = []
            step = max(1, chunk_size - overlap)
            for start in range(0, len(text), step):
                chunk = text[start:start + chunk_size].strip()
                if len(chunk) > 20:
                    chunks.append(chunk)
                if max_chunks and len(chunks) >= max_chunks:
                    logger.warning(f"Chunk limit reached ({max_chunks}). Remaining text skipped.")
                    break
            return chunks

        # Split by words for precise size control
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 для пробела
            
            # If adding word would exceed chunk size
            if current_length + word_length > chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk).strip()
                chunks.append(chunk_text)
                if max_chunks and len(chunks) >= max_chunks:
                    logger.warning(f"Chunk limit reached ({max_chunks}). Remaining text skipped.")
                    return chunks
                
                # Create overlap for next chunk
                overlap_words = []
                overlap_length = 0
                
                # Take words from end of current chunk for overlap
                for i in range(len(current_chunk) - 1, -1, -1):
                    word_len = len(current_chunk[i]) + 1
                    if overlap_length + word_len <= overlap:
                        overlap_words.insert(0, current_chunk[i])
                        overlap_length += word_len
                    else:
                        break
                
                current_chunk = overlap_words + [word]
                current_length = overlap_length + len(word)
            else:
                current_chunk.append(word)
                current_length += word_length
        
        # Add last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk).strip()
            chunks.append(chunk_text)
            if max_chunks and len(chunks) >= max_chunks:
                logger.warning(f"Достигнут лимит чанков ({max_chunks}). Остаток текста пропущен.")
                return chunks
        
        # Filter too short chunks (minimum 20 characters)
        chunks = [chunk for chunk in chunks if len(chunk.strip()) > 20]
        
        return chunks
    
    def cleanup_temp_file(self, file_path: str):
        """Удаление временного файла"""
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                logger.info(f"Temporary file deleted: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to delete temporary file {file_path}: {e}")
    
    def cleanup_all_temp_files(self):
        """Удаление всех временных файлов"""
        try:
            for file_path in self.temp_dir.iterdir():
                if file_path.is_file():
                    file_path.unlink()
            logger.info("All temporary files deleted")
        except Exception as e:
            logger.warning(f"Error cleaning temporary files: {e}")
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Получение информации о файле"""
        try:
            path = Path(file_path)
            if not path.exists():
                return {"status": "error", "error": "Файл не найден"}
            
            stat = path.stat()
            
            return {
                "filename": path.name,
                "size": stat.st_size,
                "format": path.suffix.lower(),
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "is_supported": self.is_supported_format(path.name),
                "status": "success"
            }
        except Exception as e:
            return {"status": "error", "error": str(e)}


# Global file processor instance
file_processor = FileProcessor()

def get_file_processor() -> FileProcessor:
    """Получение глобального экземпляра процессора файлов"""
    return file_processor
